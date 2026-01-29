"""
Document Loader for the Multi-Document RAG System.

Handles loading and parsing of various document formats including PDF, DOCX, and TXT.
Extracts text content while preserving metadata such as page numbers.

For DOCX files: Converts to PDF first (using docx2pdf) to extract accurate page numbers.
"""

import os
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional

from pypdf import PdfReader
from docx import Document as DocxDocument

from utils.logger import get_logger
from .document_converter import DocumentConverter

logger = get_logger(__name__)


@dataclass
class DocumentPage:
    """
    Represents a single page or section from a document.
    
    Attributes:
        document_name: Name of the source document file.
        page_number: Page number (1-indexed) or section number.
        text: The extracted text content from the page.
        is_section: True if this is a logical section (DOCX), False if actual page (PDF).
    """
    document_name: str
    page_number: int
    text: str
    is_section: bool = False  # True for DOCX sections, False for PDF pages


class DocumentLoader:
    """
    Loads and parses documents from various file formats.
    
    Supports PDF, DOCX, and TXT file formats. Extracts text content
    while preserving document structure and metadata.
    
    For DOCX files: Attempts to convert to PDF first for accurate page numbers.
    """
    
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}
    
    def __init__(self):
        """Initialize the document loader."""
        self._converter = DocumentConverter()
        logger.info("DocumentLoader initialized")
    
    def load(self, file_path: str, file_content: Optional[bytes] = None) -> List[DocumentPage]:
        """
        Load a document and extract its pages.
        
        Args:
            file_path: Path to the document file (used for extension detection and naming).
            file_content: Optional bytes content of the file. If not provided,
                         the file will be read from disk.
        
        Returns:
            List[DocumentPage]: List of pages with extracted text and metadata.
        
        Raises:
            ValueError: If the file format is not supported.
            FileNotFoundError: If file_content is None and file doesn't exist.
        """
        path = Path(file_path)
        extension = path.suffix.lower()
        document_name = path.name
        
        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file format: {extension}. "
                f"Supported formats: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )
        
        logger.info(f"Loading document: {document_name}")
        
        # For DOCX files: Try to convert to PDF first for accurate page numbers
        if extension == ".docx" and self._converter.is_available():
            # Read file content if not provided
            if file_content is None:
                with open(file_path, "rb") as f:
                    file_content = f.read()
            
            # Convert DOCX to PDF
            pdf_content = self._converter.convert_docx_to_pdf(file_content, document_name)
            
            if pdf_content:
                # Load from converted PDF - accurate page numbers!
                logger.info(f"Using converted PDF for {document_name} (accurate page numbers)")
                return self._load_pdf(document_name, file_path, pdf_content)
            else:
                # Conversion failed, fall back to DOCX parsing
                logger.warning(f"PDF conversion failed for {document_name}, using section-based parsing")
        
        if extension == ".pdf":
            return self._load_pdf(document_name, file_path, file_content)
        elif extension == ".docx":
            return self._load_docx(document_name, file_path, file_content)
        elif extension == ".txt":
            return self._load_txt(document_name, file_path, file_content)
        
        return []
    
    def _load_pdf(
        self, 
        document_name: str, 
        file_path: str, 
        file_content: Optional[bytes]
    ) -> List[DocumentPage]:
        """
        Load and parse a PDF document.
        
        Args:
            document_name: Name of the document file.
            file_path: Path to the PDF file.
            file_content: Optional bytes content of the file.
        
        Returns:
            List[DocumentPage]: Extracted pages with text content.
        """
        pages = []
        
        try:
            if file_content:
                import io
                reader = PdfReader(io.BytesIO(file_content))
            else:
                reader = PdfReader(file_path)
            
            for page_num, page in enumerate(reader.pages, start=1):
                text = page.extract_text() or ""
                text = text.strip()
                
                if text:  # Only add pages with content
                    pages.append(DocumentPage(
                        document_name=document_name,
                        page_number=page_num,
                        text=text
                    ))
            
            logger.info(f"Extracted {len(pages)} pages from PDF: {document_name}")
            
        except Exception as e:
            logger.error(f"Error loading PDF {document_name}: {str(e)}")
            raise
        
        return pages
    
    def _load_docx(
        self, 
        document_name: str, 
        file_path: str, 
        file_content: Optional[bytes]
    ) -> List[DocumentPage]:
        """
        Load and parse a DOCX document.
        
        Extracts text from paragraphs, tables, headers, and footers.
        DOCX doesn't have native page info, so content is split into
        logical sections or returned as a single page.
        
        Args:
            document_name: Name of the document file.
            file_path: Path to the DOCX file.
            file_content: Optional bytes content of the file.
        
        Returns:
            List[DocumentPage]: Extracted pages with text content.
        """
        pages = []
        
        try:
            import io
            if file_content:
                doc = DocxDocument(io.BytesIO(file_content))
            else:
                doc = DocxDocument(file_path)
            
            all_text_parts = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    all_text_parts.append(text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        all_text_parts.append(" | ".join(row_text))
            
            # Extract text from headers (if any)
            for section in doc.sections:
                header = section.header
                if header:
                    for para in header.paragraphs:
                        text = para.text.strip()
                        if text:
                            all_text_parts.insert(0, f"[Header] {text}")
                
                footer = section.footer
                if footer:
                    for para in footer.paragraphs:
                        text = para.text.strip()
                        if text:
                            all_text_parts.append(f"[Footer] {text}")
            
            # Combine all text
            full_text = "\n\n".join(all_text_parts)
            
            if full_text:
                # If document is very long, split into multiple "sections"
                max_page_chars = 3000
                if len(full_text) > max_page_chars:
                    # Split into logical chunks
                    chunks = self._split_into_pages(full_text, max_page_chars)
                    for i, chunk in enumerate(chunks, start=1):
                        pages.append(DocumentPage(
                            document_name=document_name,
                            page_number=i,
                            text=chunk,
                            is_section=True  # Mark as section, not actual page
                        ))
                else:
                    pages.append(DocumentPage(
                        document_name=document_name,
                        page_number=1,
                        text=full_text,
                        is_section=True  # Mark as section, not actual page
                    ))
            
            logger.info(f"Extracted {len(pages)} page(s) from DOCX: {document_name}")
            
        except Exception as e:
            logger.error(f"Error loading DOCX {document_name}: {str(e)}")
            raise
        
        return pages
    
    def _split_into_pages(self, text: str, max_chars: int) -> List[str]:
        """
        Split text into page-like chunks.
        
        Args:
            text: Full document text.
            max_chars: Maximum characters per page.
        
        Returns:
            List[str]: List of text chunks.
        """
        chunks = []
        paragraphs = text.split("\n\n")
        current_chunk = []
        current_length = 0
        
        for para in paragraphs:
            if current_length + len(para) > max_chars and current_chunk:
                chunks.append("\n\n".join(current_chunk))
                current_chunk = [para]
                current_length = len(para)
            else:
                current_chunk.append(para)
                current_length += len(para)
        
        if current_chunk:
            chunks.append("\n\n".join(current_chunk))
        
        return chunks

    
    def _load_txt(
        self, 
        document_name: str, 
        file_path: str, 
        file_content: Optional[bytes]
    ) -> List[DocumentPage]:
        """
        Load and parse a plain text document.
        
        Args:
            document_name: Name of the document file.
            file_path: Path to the TXT file.
            file_content: Optional bytes content of the file.
        
        Returns:
            List[DocumentPage]: Single page containing all text content.
        """
        pages = []
        
        try:
            if file_content:
                text = file_content.decode("utf-8")
            else:
                with open(file_path, "r", encoding="utf-8") as f:
                    text = f.read()
            
            text = text.strip()
            
            if text:
                pages.append(DocumentPage(
                    document_name=document_name,
                    page_number=1,
                    text=text
                ))
            
            logger.info(f"Extracted content from TXT: {document_name}")
            
        except Exception as e:
            logger.error(f"Error loading TXT {document_name}: {str(e)}")
            raise
        
        return pages
    
    @classmethod
    def is_supported(cls, file_path: str) -> bool:
        """
        Check if a file format is supported.
        
        Args:
            file_path: Path to the file to check.
        
        Returns:
            bool: True if the file format is supported.
        """
        extension = Path(file_path).suffix.lower()
        return extension in cls.SUPPORTED_EXTENSIONS
